import os
import re
import mammoth
from docx import Document
from docx.shared import Pt, Inches
from docx2pdf import convert
from pypdf import PdfReader

class FolderCampaignLoader:
    def __init__(self, folder_path: str):
        self.folder_path = folder_path
        self.subject_file = os.path.join(folder_path, "Subject.txt")
        self.email_file = os.path.join(folder_path, "Email.docx")
        
        cv_docx = os.path.join(folder_path, "CV_FAC.docx")
        cv_pdf = os.path.join(folder_path, "CV_FAC.pdf")
        
        if os.path.exists(cv_docx):
            self.cv_file = cv_docx
            self.cv_type = "docx"
        elif os.path.exists(cv_pdf):
            self.cv_file = cv_pdf
            self.cv_type = "pdf"
        else:
            self.cv_file = None
            self.cv_type = None

    def validate_folder(self) -> tuple[bool, str]:
        if not os.path.exists(self.folder_path):
            return False, "Folder path does not exist."
        if not os.path.exists(self.subject_file):
            return False, "Missing 'Subject.txt' in folder."
        if not os.path.exists(self.email_file):
            return False, "Missing 'Email.docx' in folder."
        if not self.cv_file:
            return False, "Missing 'CV_FAC.docx' or 'CV_FAC.pdf' in folder."
        return True, "All campaign assets found!"

    def load_campaign(self) -> dict:
        valid, msg = self.validate_folder()
        if not valid:
            raise FileNotFoundError(msg)

        # 1. Read Subject.txt
        with open(self.subject_file, "r", encoding="utf-8") as f:
            subject = f.read().strip()

        # 2. Convert Email.docx (including inline figures) to HTML.
        # Gmail strips base64 data-URI <img> tags on display, so instead we
        # give each image a Content-ID and reference it with cid:, then attach
        # the actual image bytes inline (done in gmail_sender). This is what
        # makes inline figures survive in received Gmail.
        inline_images = []  # list of dicts: {cid, data, content_type}

        def _convert_image(image):
            index = len(inline_images)
            cid = f"img{index}"
            with image.open() as image_bytes:
                data = image_bytes.read()
            inline_images.append({
                "cid": cid,
                "data": data,
                "content_type": image.content_type or "image/png",
            })
            return {"src": f"cid:{cid}"}

        with open(self.email_file, "rb") as docx_file:
            result = mammoth.convert_to_html(
                docx_file,
                convert_image=mammoth.images.img_element(_convert_image),
            )
            html_body = result.value

        return {
            "subject": subject,
            "html_body": html_body,
            "cv_path": self.cv_file,
            "cv_type": self.cv_type,
            "inline_images": inline_images,
        }

    @staticmethod
    def _adjust_docx_spacing(doc: Document, tight_level: int = 1):
        """
        Dynamically adjusts margins and line/paragraph spacing to keep CV layout professional.
        """
        margin_val = max(0.45, 0.65 - (0.08 * tight_level))
        for section in doc.sections:
            section.top_margin = Inches(margin_val)
            section.bottom_margin = Inches(margin_val)
            section.left_margin = Inches(0.55)
            section.right_margin = Inches(0.55)

        for p in doc.paragraphs:
            p.paragraph_format.line_spacing = max(1.0, 1.12 - (0.04 * tight_level))
            if p.paragraph_format.space_after and p.paragraph_format.space_after > Pt(2):
                p.paragraph_format.space_after = Pt(max(1, 4 - tight_level))
            if p.paragraph_format.space_before and p.paragraph_format.space_before > Pt(2):
                p.paragraph_format.space_before = Pt(max(1, 4 - tight_level))

    @staticmethod
    def generate_tailored_cv(docx_path: str, output_pdf_path: str, replacements: dict) -> str:
        """
        Loads CV_FAC.docx, applies replacements, and optimizes PDF layout to target 
        5 exact pages (or 6 full pages max) without half-empty final pages.
        """
        doc = Document(docx_path)

        for p in doc.paragraphs:
            for key, val in replacements.items():
                if key in p.text:
                    p.text = p.text.replace(key, val)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for key, val in replacements.items():
                            if key in p.text:
                                p.text = p.text.replace(key, val)

        temp_docx = output_pdf_path.replace(".pdf", "_temp.docx")

        for attempt in range(4):
            doc.save(temp_docx)
            convert(temp_docx, output_pdf_path)

            try:
                reader = PdfReader(output_pdf_path)
                num_pages = len(reader.pages)
                last_page_idx = num_pages - 1
                last_page_chars = len(reader.pages[last_page_idx].extract_text().strip())

                if num_pages > 6:
                    FolderCampaignLoader._adjust_docx_spacing(doc, tight_level=attempt + 1)
                elif num_pages in [5, 6]:
                    if last_page_chars < 600:
                        FolderCampaignLoader._adjust_docx_spacing(doc, tight_level=attempt + 1)
                    else:
                        break
                else:
                    break
            except Exception as e:
                print(f"PDF layout inspection note: {e}")
                break

        if os.path.exists(temp_docx):
            os.remove(temp_docx)

        return output_pdf_path

    @staticmethod
    def personalize_text(text: str, first_name: str, last_name: str, title: str, university: str = "", profile_match: str = "") -> str:
        salutation_title = title if title else "Professor"
        display_last_name = last_name if last_name else "Professor"

        personalized = re.sub(
            r"Dear\s+(Dr\.|Prof\.|Professor)\s*,",
            f"Dear \\1 {display_last_name},",
            text,
            flags=re.IGNORECASE
        )

        personalized = personalized.replace("{FIRST_NAME}", first_name)
        personalized = personalized.replace("{LAST_NAME}", display_last_name)
        personalized = personalized.replace("{TITLE}", salutation_title)
        personalized = personalized.replace("{UNIVERSITY}", university)
        
        if "{PROFILE_MATCH}" in personalized:
            replacement = f" {profile_match}" if profile_match else ""
            personalized = personalized.replace("{PROFILE_MATCH}", replacement)

        return personalized